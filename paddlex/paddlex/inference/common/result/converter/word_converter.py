# Copyright (c) 2024 PaddlePaddle Authors. All Rights Reserved.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#    http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""WordConverter — converts structured word_blocks to a docx.Document."""

from __future__ import annotations

import math
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

# Block type labels that represent image-like content (chart/image/seal)
_IMAGE_LABELS = ("chart", "image", "seal")

# Maximum rendered height for header/footer images (keeps logos from expanding the band)
_MAX_HEADER_IMG_HEIGHT_EMU = 457200  # 0.5 inch in EMU

# Regex to detect $$ display $$ or $ inline $ formula markers in plain text
_INLINE_FORMULA_RE = re.compile(r"(\$\$[\s\S]*?\$\$|\$[^$\n]+?\$)")

# Formula block labels
_FORMULA_LABELS = ("inline_formula", "display_formula", "formula")

# Cached XSLT transform for LaTeX→OMML conversion (lazy-loaded on first use)
_OMML_TRANSFORM = None


def _get_omml_transform():
    """Return (and cache) the XSLT transform object for MathML→OMML conversion."""
    global _OMML_TRANSFORM
    if _OMML_TRANSFORM is None:
        from lxml import etree as _etree

        xsl_path = Path(__file__).parent / "MML2OMML.XSL"
        _OMML_TRANSFORM = _etree.XSLT(_etree.parse(str(xsl_path)))
    return _OMML_TRANSFORM


def _split_inline_formulas(text: str) -> List[Tuple[str, bool]]:
    """Split text into (segment, is_formula) pairs.

    is_formula=True if the segment is a $...$ or $$...$$ formula marker.
    """
    parts: List[Tuple[str, bool]] = []
    last_end = 0
    for m in _INLINE_FORMULA_RE.finditer(text):
        if m.start() > last_end:
            parts.append((text[last_end : m.start()], False))
        parts.append((m.group(), True))
        last_end = m.end()
    if last_end < len(text):
        parts.append((text[last_end:], False))
    return parts


def _strip_latex_markers(content: str) -> Tuple[str, bool]:
    """Strip $/$$ markers from formula content.

    Returns:
        (raw_latex, is_display) where is_display=True for $$ or \\[...\\].
    """
    s = content.strip()
    if s.startswith("$$") and s.endswith("$$"):
        return s[2:-2].strip(), True
    if s.startswith("\\[") and s.endswith("\\]"):
        return s[2:-2].strip(), True
    if s.startswith("$") and s.endswith("$"):
        return s[1:-1].strip(), False
    if s.startswith("\\(") and s.endswith("\\)"):
        return s[2:-2].strip(), False
    # No markers → treat as display formula (most formula blocks are display)
    return s, True


_OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_OMML_SCRIPT_TAGS = ("sSup", "sSub", "sSubSup")
_ZWSP = "\u200b"  # zero-width space — invisible placeholder, prevents Word □ box


def _fill_empty_bases(omml_root):
    """Replace empty <m:e/> in sSup/sSub/sSubSup with a zero-width space run.

    Word renders a visible placeholder box (□) for truly empty <m:e/> elements.
    Inserting a zero-width character suppresses the box while remaining invisible.
    """
    from lxml import etree as _etree

    ns = _OMML_NS
    for tag in _OMML_SCRIPT_TAGS:
        for base_elem in omml_root.findall(f".//{{{ns}}}{tag}/{{{ns}}}e"):
            if len(base_elem) == 0 and not (base_elem.text or "").strip():
                mr = _etree.SubElement(base_elem, f"{{{ns}}}r")
                mt = _etree.SubElement(mr, f"{{{ns}}}t")
                mt.text = _ZWSP
    return omml_root


def _latex_to_omml(latex_str: str, display: bool = False):
    """Convert a LaTeX string to an OMML XML element (<m:oMath>).

    Args:
        latex_str: Raw LaTeX (without surrounding $ markers).
        display: True for display-mode (block), False for inline.

    Returns:
        lxml Element (<m:oMath>) or None if conversion fails.

    Raises:
        ImportError: If latex2mathml or lxml is not installed.
    """
    import latex2mathml.converter
    from lxml import etree as _etree

    try:
        # Normalize bare ^ or _ (no base) to {}^{} or {}_{} per LaTeX spec
        if latex_str.startswith(("^", "_")):
            latex_str = "{}" + latex_str
        mode = "block" if display else "inline"
        mathml = latex2mathml.converter.convert(latex_str, display=mode)
        transform = _get_omml_transform()
        mml_root = _etree.fromstring(mathml.encode())
        omml_tree = transform(mml_root)
        return _fill_empty_bases(omml_tree.getroot())  # <m:oMath> element
    except Exception:
        return None


# Labels excluded from body content (written to section header/footer or skipped)
_HEADER_FOOTER_LABELS = {
    "header",
    "footer",
    "header_image",
    "footer_image",
    "aside_text",
}


def _get_image_size(abs_path: str) -> Optional[Tuple[int, int]]:
    """Return (width, height) in pixels for the image at abs_path, or None on error."""
    try:
        from PIL import Image as _PILImage

        img = _PILImage.open(abs_path)
        size = img.size
        img.close()
        return size
    except Exception:
        return None


def _header_image_width(
    abs_path: str,
    bbox,
    original_image_width: int,
    usable_width_emu: Optional[int],
) -> int:
    """Compute display width (EMU) for a header/footer image.

    Scales proportionally from bbox, capped by usable page width and
    _MAX_HEADER_IMG_HEIGHT_EMU (aspect-ratio preserving).
    Falls back to Inches(1.0) when bbox or image dimensions are unavailable.
    """
    from docx.shared import Inches

    if bbox and original_image_width > 0 and usable_width_emu:
        ratio = (bbox[2] - bbox[0]) / original_image_width
        width = max(Inches(0.2), min(int(ratio * usable_width_emu), usable_width_emu))
        dims = _get_image_size(abs_path)
        if dims:
            natural_w, natural_h = dims
            if natural_w > 0:
                rendered_h = int(width * natural_h / natural_w)
                if rendered_h > _MAX_HEADER_IMG_HEIGHT_EMU:
                    width = max(
                        Inches(0.2),
                        int(_MAX_HEADER_IMG_HEIGHT_EMU * natural_w / natural_h),
                    )
        return width
    return Inches(1.0)  # fallback


def _set_paragraph_style(para, config):
    """Apply font/alignment config to a Word paragraph."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    run = para.runs[0] if para.runs else para.add_run()
    font_name = config.get("font", "Times New Roman")
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(config.get("size", 12))
    run.bold = config.get("bold", False)
    para.alignment = config.get("align", WD_ALIGN_PARAGRAPH.LEFT)
    if config.get("indent", False):
        para.paragraph_format.first_line_indent = Inches(0.3)
    # Force single line spacing to prevent default 1.15x from consuming extra vertical space
    para.paragraph_format.line_spacing = 1.0


def _write_mixed_runs(para, parts: List[Tuple[str, bool]], config: dict) -> None:
    """Write alternating text/formula segments into an existing paragraph.

    For formula segments: attempts OMML conversion; falls back to plain text.
    For text segments: adds a styled run with font settings from config.
    """
    from docx.oxml.ns import qn
    from docx.shared import Pt

    font_name = config.get("font", "Times New Roman")
    font_size = config.get("size", 12)
    bold = config.get("bold", False)

    for segment, is_formula in parts:
        if not segment:
            continue
        if is_formula:
            raw_latex, is_display = _strip_latex_markers(segment)
            omml_elem = (
                _latex_to_omml(raw_latex, display=is_display) if raw_latex else None
            )
            if omml_elem is not None:
                para._element.append(omml_elem)
                continue
            # Fallback: write as plain text run
        run = para.add_run(segment)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(font_size)
        run.bold = bold


def _classify_number_position(bbox, page_width, page_height):
    """Classify a 'number' block's semantic role based on its bbox position.

    Args:
        bbox: [x1, y1, x2, y2] bounding box in pixel coordinates.
        page_width: Page width in pixels.
        page_height: Page height in pixels.

    Returns:
        One of: 'header', 'footer', 'aside_text'.
    """
    if not bbox or page_width <= 0 or page_height <= 0:
        return "footer"

    x1, y1, x2, y2 = bbox
    y_center = (y1 + y2) / 2.0
    x_center = (x1 + x2) / 2.0

    # Top 10% → header region
    if y_center < page_height * 0.10:
        return "header"

    # Bottom 10% → footer region
    if y_center > page_height * 0.90:
        return "footer"

    # Left 15% or right 15% (not in header/footer zone) → aside_text
    if x_center < page_width * 0.15 or x_center > page_width * 0.85:
        return "aside_text"

    # Default fallback
    return "footer"


def _parse_html_table(html: str) -> List[List[List[Tuple[str, str]]]]:
    """Parse an HTML table into rows of cells, each cell a list of (kind, value) segments.

    kind is "text" for text content or "img" for image src paths.
    """
    from bs4 import BeautifulSoup, Tag
    from bs4.element import NavigableString

    soup = BeautifulSoup(html, "html.parser")
    rows = []
    for tr in soup.find_all("tr"):
        if not isinstance(tr, Tag):
            continue
        cells = []
        for cell in tr.find_all(["td", "th"]):
            if not isinstance(cell, Tag):
                continue
            segments: List[Tuple[str, str]] = []
            for child in cell.children:
                if isinstance(child, NavigableString):
                    text = child.strip()
                    if text:
                        segments.append(("text", text))
                elif isinstance(child, Tag) and child.name == "img":
                    src = str(child.get("src", "") or "")
                    if src:
                        segments.append(("img", src))
                elif isinstance(child, Tag):
                    text = child.get_text(strip=True)
                    if text:
                        segments.append(("text", text))
            cells.append(segments)
        rows.append(cells)
    return rows


def build_word_blocks(
    parsing_res_list: List[Any],
    extra_style_map: Optional[Dict[str, Dict]] = None,
    page_width: int = 0,
    page_height: int = 0,
    imgs_in_doc: Optional[List[Dict]] = None,
) -> tuple:
    """Build word_blocks and images list from a parsing_res_list.

    Extracts the shared logic for converting DocumentBlock / PaddleOCRVLBlock
    objects into the word_blocks format expected by WordConverter.convert().

    Args:
        parsing_res_list: List of block objects with .label, .content, .image attrs.
        extra_style_map: Optional dict of label->style overrides merged on top of
            BASE_STYLE_MAP via dict.update(). Use for pipeline-specific labels.
        page_width: Page width in pixels, used to classify 'number' blocks.
            0 means unknown (defaults to footer classification).
        page_height: Page height in pixels, used to classify 'number' blocks.
            0 means unknown (defaults to footer classification).
        imgs_in_doc: Optional list of {"path": str, "img": PIL.Image} dicts from
            self["imgs_in_doc"], covering images embedded in table cells. These
            are merged into the returned images list (parallel to MarkdownConverter).

    Returns:
        Tuple of (word_blocks, images) where:
            word_blocks: List[Dict] with keys "type", "content", "config",
                and optionally "bbox" and "page_index".
            images: List[Dict] with keys "path" and "img".
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    BASE_STYLE_MAP = {
        "doc_title": {
            "level": 0,
            "size": 20,
            "bold": True,
            "align": WD_ALIGN_PARAGRAPH.CENTER,
        },
        "header": {
            "size": 16,
            "bold": True,
            "align": WD_ALIGN_PARAGRAPH.CENTER,
        },
        "abstract_title": {
            "level": 1,
            "size": 14,
            "bold": True,
            "align": WD_ALIGN_PARAGRAPH.CENTER,
        },
        "content_title": {
            "level": 1,
            "size": 14,
            "bold": True,
            "align": WD_ALIGN_PARAGRAPH.LEFT,
        },
        "reference_title": {
            "level": 1,
            "size": 14,
            "bold": True,
            "align": WD_ALIGN_PARAGRAPH.LEFT,
        },
        "paragraph_title": {
            "level": 2,
            "size": 14,
            "bold": True,
            "align": WD_ALIGN_PARAGRAPH.LEFT,
        },
        "abstract": {"size": 12, "align": WD_ALIGN_PARAGRAPH.JUSTIFY},
        "text": {
            "size": 12,
            "align": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "indent": True,
        },
        "figure_title": {"size": 10, "align": WD_ALIGN_PARAGRAPH.CENTER},
        "table_title": {"size": 10, "align": WD_ALIGN_PARAGRAPH.CENTER},
        "chart_title": {"size": 10, "align": WD_ALIGN_PARAGRAPH.CENTER},
        "reference": {"size": 12, "align": WD_ALIGN_PARAGRAPH.JUSTIFY},
        "algorithm": {
            "font": "Courier New",
            "size": 11,
            "align": WD_ALIGN_PARAGRAPH.LEFT,
        },
        "formula": {"size": 12, "align": WD_ALIGN_PARAGRAPH.CENTER},
        "vision_footnote": {"size": 9, "align": WD_ALIGN_PARAGRAPH.LEFT},
        "number": {"size": 9, "align": WD_ALIGN_PARAGRAPH.CENTER},
        "footer": {"size": 9, "align": WD_ALIGN_PARAGRAPH.CENTER},
    }

    style_map = {**BASE_STYLE_MAP}
    if extra_style_map:
        style_map.update(extra_style_map)

    default_config = {"size": 12, "align": WD_ALIGN_PARAGRAPH.LEFT, "indent": True}

    word_blocks = []
    images = []

    for block in parsing_res_list:
        label = block.label
        content = getattr(block, "content", "")
        if label in ["image", "seal"]:
            if block.image is None:
                continue
            content = block.image["path"]
        elif label in ("header_image", "footer_image"):
            if block.image is None:
                continue
            content = block.image["path"]
        elif label == "chart":
            if block.image is not None:
                content = block.image["path"]
            elif content:
                # VLM chart recognition: pipe-delimited table text → reuse table rendering
                content = content.replace("|", "\t")
                label = "table"
            else:
                continue
        elif label == "number":
            # Classify 'number' blocks by position to reuse header/footer/aside_text paths
            bbox = (
                list(block.bbox)
                if hasattr(block, "bbox") and block.bbox is not None
                else None
            )
            label = _classify_number_position(
                bbox, page_width=page_width, page_height=page_height
            )
        config = style_map.get(label, default_config)
        word_block = {
            "type": label,
            "content": content,
            "config": config,
        }
        if hasattr(block, "bbox") and block.bbox is not None:
            word_block["bbox"] = list(block.bbox)
        if hasattr(block, "page_index") and block.page_index is not None:
            word_block["page_index"] = block.page_index
        word_blocks.append(word_block)
        if block.image is not None:
            images.append({"path": block.image["path"], "img": block.image["img"]})

    # Include table-embedded images (parallel to MarkdownConverter.convert() logic)
    if imgs_in_doc:
        existing_paths = {img["path"] for img in images}
        for item in imgs_in_doc:
            if (
                item.get("path")
                and item.get("img") is not None
                and item["path"] not in existing_paths
            ):
                images.append({"path": item["path"], "img": item["img"]})

    return word_blocks, images


def _write_block(
    doc,
    block,
    abs_image_paths,
    original_image_width=500,
    space_before_emu=None,
    left_indent_emu=None,
    usable_width_emu=None,
    max_height_emu=None,
):
    """Write a single word_block to the given docx Document (or container).

    Handles image/chart/seal, table, and text blocks. Header/footer blocks
    are intentionally NOT handled here — callers must write them to
    section.header / section.footer separately.

    Args:
        doc: docx.Document or a document-like container supporting
            add_paragraph() / add_table().
        block: Dict with keys "type", "content", "config".
        abs_image_paths: Dict mapping original image path → absolute path.
        original_image_width: Width of the original page image in pixels, used to
            calculate proportional image width in the Word document.
        space_before_emu: Optional space before this block in EMU.
        left_indent_emu: Optional left indent in EMU (single-column only).
        usable_width_emu: Optional usable page width in EMU for proportional sizing.
        max_height_emu: Optional maximum rendered height in EMU for image scaling.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Emu, Inches, Pt

    label = block.get("type")
    content = block.get("content", "")
    if isinstance(content, str):
        content = content.strip()
    config = block.get("config") or {}

    # --- image/chart/seal ---
    if label in _IMAGE_LABELS:
        image_name = block.get("content")
        if not image_name:
            return
        abs_image_path = abs_image_paths.get(image_name)
        if not abs_image_path:
            return
        para = doc.add_paragraph()
        if space_before_emu is not None:
            para.paragraph_format.space_before = Emu(space_before_emu)
            para.paragraph_format.space_after = Emu(0)
        run = para.add_run()
        # Calculate proportional width based on bbox ratio
        USABLE_PAGE_WIDTH = 6.0  # inches fallback
        bbox = block.get("bbox")
        if bbox and original_image_width > 0:
            ratio = (bbox[2] - bbox[0]) / original_image_width
            if usable_width_emu:
                img_width = max(
                    Inches(0.1), min(int(ratio * usable_width_emu), usable_width_emu)
                )
                # Apply max_height_emu constraint (aspect-ratio preserving)
                if max_height_emu and max_height_emu > 0:
                    dims = _get_image_size(abs_image_path)
                    if dims:
                        natural_w, natural_h = dims
                        rendered_h = int(img_width * natural_h / natural_w)
                        if rendered_h > max_height_emu:
                            img_width = int(max_height_emu * natural_w / natural_h)
                            img_width = max(Inches(0.1), img_width)
                run.add_picture(abs_image_path, width=img_width)
            else:
                img_width = max(0.1, min(ratio * USABLE_PAGE_WIDTH, USABLE_PAGE_WIDTH))
                run.add_picture(abs_image_path, width=Inches(img_width))
        else:
            run.add_picture(abs_image_path, width=Inches(5.0))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- table ---
    elif label == "table" and content:
        if "<table" in content:
            rows = _parse_html_table(content)
        else:
            # Plain-text table: wrap each cell string as a single text segment
            rows = [
                [[("text", c)] for c in r.split("\t")]
                for r in content.split("\n")
                if r.strip()
            ]
        if rows:
            # Insert spacer paragraph for spacing before table
            if space_before_emu is not None and space_before_emu > 0:
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_before = Emu(space_before_emu)
                spacer.paragraph_format.space_after = Emu(0)
                spacer.paragraph_format.line_spacing = Pt(1)
                run = spacer.add_run()
                run.font.size = Pt(1)

            max_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=0, cols=max_cols)
            table.style = "Table Grid"

            # Set proportional table width from bbox; also compute col_width for images
            col_width = None
            bbox = block.get("bbox")
            if bbox and original_image_width > 0 and usable_width_emu:
                ratio = (bbox[2] - bbox[0]) / original_image_width
                table_width = max(Inches(2), int(ratio * usable_width_emu))
                col_width = table_width // max_cols
                for col in table.columns:
                    col.width = col_width

            for row_cells in rows:
                row = table.add_row().cells
                for i in range(max_cols):
                    segments = row_cells[i] if i < len(row_cells) else [("text", "")]
                    cell_para = row[i].paragraphs[0]
                    for seg_kind, seg_val in segments:
                        if seg_kind == "text":
                            text = seg_val.strip()
                            if not text:
                                continue
                            if "$" in text:
                                parts = _split_inline_formulas(text)
                                _write_mixed_runs(cell_para, parts, {})
                            else:
                                run = cell_para.add_run(text)
                                run.font.name = "Times New Roman"
                                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                                run.font.size = Pt(12)
                        elif seg_kind == "img":
                            abs_path = abs_image_paths.get(seg_val)
                            if abs_path:
                                img_w = (
                                    int(col_width * 0.9) if col_width else Inches(1.0)
                                )
                                cell_para.add_run().add_picture(abs_path, width=img_w)

    # --- formula (inline_formula / display_formula / formula) ---
    elif label in _FORMULA_LABELS and content:
        raw_latex, is_display = _strip_latex_markers(content)
        omml_elem = _latex_to_omml(raw_latex, display=is_display) if raw_latex else None

        para = doc.add_paragraph()
        _set_paragraph_style(para, config)
        if space_before_emu is not None:
            para.paragraph_format.space_before = Emu(space_before_emu)
            para.paragraph_format.space_after = Emu(0)
        if left_indent_emu is not None:
            para.paragraph_format.left_indent = Emu(left_indent_emu)

        if omml_elem is not None:
            para._element.append(omml_elem)
        else:
            # Fallback: write raw content as plain text
            para.add_run(content)

    # --- other text content (including text blocks with possible $...$ formulas) ---
    elif (
        label
        not in [
            "header",
            "footer",
            "header_image",
            "footer_image",
            "table",
            "chart",
            "image",
            "seal",
            "aside_text",
        ]
        and content
    ):
        lines = [l for l in content.split("\n") if l.strip()]
        for i, line in enumerate(lines):
            first = i == 0
            if "$" in line:
                parts = _split_inline_formulas(line)
                has_formula = any(is_formula for _, is_formula in parts)
            else:
                parts = []
                has_formula = False

            if has_formula:
                para = doc.add_paragraph()
                _set_paragraph_style(para, config)
                if first and space_before_emu is not None:
                    para.paragraph_format.space_before = Emu(space_before_emu)
                    para.paragraph_format.space_after = Emu(0)
                if left_indent_emu is not None:
                    para.paragraph_format.left_indent = Emu(left_indent_emu)
                _write_mixed_runs(para, parts, config)
            else:
                para = doc.add_paragraph(line)
                _set_paragraph_style(para, config)
                if first and space_before_emu is not None:
                    para.paragraph_format.space_before = Emu(space_before_emu)
                    para.paragraph_format.space_after = Emu(0)
                if left_indent_emu is not None:
                    para.paragraph_format.left_indent = Emu(left_indent_emu)


def _build_page_metrics(body_blocks, page_width_px, page_height_px):
    """Compute layout metrics for one page.

    Args:
        body_blocks: List of block dicts (header/footer/aside_text already excluded).
        page_width_px: Original page width in pixels.
        page_height_px: Original page height in pixels.

    Returns:
        Dict with keys:
            scale_x, scale_y: float (px to EMU)
            content_bbox: (x1, y1, x2, y2) in px — body content bounding box
            margins: (left, right, top, bottom) in EMU
            usable_width_emu: int — page usable width after margins
    """
    # A4: 210mm x 297mm = 7560820 x 10693400 EMU
    PAGE_WIDTH_EMU = 7560820
    PAGE_HEIGHT_EMU = 10693400
    MIN_MARGIN_EMU = 274320  # 0.3 inch
    MAX_MARGIN_EMU = 1828800  # 2.0 inch

    scale_x = PAGE_WIDTH_EMU / max(page_width_px, 1)
    scale_y = PAGE_HEIGHT_EMU / max(page_height_px, 1)

    blocks_with_bbox = [b for b in body_blocks if b.get("bbox")]
    if not blocks_with_bbox:
        # Default margins: 1 inch on all sides
        default_margin = 914400
        usable = PAGE_WIDTH_EMU - 2 * default_margin
        usable_h = PAGE_HEIGHT_EMU - 2 * default_margin
        return {
            "scale_x": scale_x,
            "scale_y": scale_y,
            "content_bbox": None,
            "margins": (default_margin, default_margin, default_margin, default_margin),
            "usable_width_emu": usable,
            "usable_height_emu": max(usable_h, 1),
        }

    x1s, y1s, x2s, y2s = zip(*(b["bbox"][:4] for b in blocks_with_bbox))

    content_x1, content_y1 = min(x1s), min(y1s)
    content_x2, content_y2 = max(x2s), max(y2s)

    left_px = content_x1
    right_px = max(0, page_width_px - content_x2)
    top_px = content_y1
    bottom_px = max(0, page_height_px - content_y2)

    def _clamp(val_emu):
        return max(MIN_MARGIN_EMU, min(MAX_MARGIN_EMU, val_emu))

    left_m = _clamp(int(left_px * scale_x))
    right_m = _clamp(int(right_px * scale_x))
    top_m = _clamp(int(top_px * scale_y))
    bottom_m = _clamp(int(bottom_px * scale_y))

    usable_width_emu = PAGE_WIDTH_EMU - left_m - right_m
    usable_height_emu = PAGE_HEIGHT_EMU - top_m - bottom_m

    return {
        "scale_x": scale_x,
        "scale_y": scale_y,
        "content_bbox": (content_x1, content_y1, content_x2, content_y2),
        "margins": (left_m, right_m, top_m, bottom_m),
        "usable_width_emu": max(usable_width_emu, 1),
        "usable_height_emu": max(usable_height_emu, 1),
    }


def _compute_vertical_spacing(blocks, scale_y):
    """Compute space_before (EMU) for each block based on y-gap from previous block.

    Args:
        blocks: List of block dicts with "bbox" key, sorted by y.
        scale_y: Pixels-to-EMU conversion factor for Y axis.

    Returns:
        List of int|None, same length as blocks. None means use default spacing.
        First block always returns 0.
    """
    MAX_SPACE_EMU = 914400  # 1 inch cap
    QUANTIZE_STEP = 38100  # 3pt quantization to reduce OCR bbox noise

    spacings = []
    prev_y2 = None
    for block in blocks:
        bbox = block.get("bbox")
        if not bbox:
            spacings.append(None)
            prev_y2 = None
            continue
        y1, y2 = bbox[1], bbox[3]
        if prev_y2 is None:
            spacings.append(0)
        else:
            gap_px = max(0, y1 - prev_y2)
            space_emu = int(gap_px * scale_y)
            space_emu = min(space_emu, MAX_SPACE_EMU)
            # Quantize to 3pt steps to reduce OCR noise
            space_emu = round(space_emu / QUANTIZE_STEP) * QUANTIZE_STEP
            spacings.append(space_emu)
        prev_y2 = y2
    return spacings


def _estimate_block_height(
    block, column_width_emu, abs_image_paths, scale_x, scale_y, original_image_width=0
):
    """Estimate the rendered height of a single block in Word (EMU).

    Args:
        block: Block dict with "type", "content", "config", optional "bbox".
        column_width_emu: Available column width in EMU.
        abs_image_paths: Dict mapping image name to absolute path.
        scale_x: Pixels-to-EMU X factor.
        scale_y: Pixels-to-EMU Y factor.
        original_image_width: Original page width in pixels, used for image
            aspect-ratio calculation (same value as passed to _write_block).

    Returns:
        int: Estimated height in EMU.
    """
    from docx.shared import Inches

    label = block.get("type", "")
    bbox = block.get("bbox")
    config = block.get("config") or {}
    content = block.get("content", "")
    if isinstance(content, str):
        content = content.strip()

    LINE_HEIGHT_FACTOR = 1.2  # Word line height ≈ font_size × 1.2

    if label in _IMAGE_LABELS:
        image_name = block.get("content")
        abs_path = abs_image_paths.get(image_name) if image_name else None
        if abs_path and bbox and column_width_emu > 0:
            dims = _get_image_size(abs_path)
            if dims:
                natural_w, natural_h = dims
                # Use original_image_width (page pixel width) as denominator — same
                # as _write_block — so the ratio is a fraction of the page, not column.
                ref_width = (
                    original_image_width
                    if original_image_width > 0
                    else max(1, int(column_width_emu / scale_x))
                )
                ratio = (bbox[2] - bbox[0]) / max(ref_width, 1)
                img_width = max(
                    Inches(1.0), min(int(ratio * column_width_emu), column_width_emu)
                )
                rendered_h = int(img_width * natural_h / max(natural_w, 1))
                return max(rendered_h, 914400 // 10)  # min 0.1"
        # Fallback: bbox-based
        if bbox:
            return int((bbox[3] - bbox[1]) * scale_y)
        return int(Inches(2.0))  # type: ignore[return-value]

    if label == "table":
        if bbox:
            bbox_h_px = bbox[3] - bbox[1]
            bbox_w_px = max(1, bbox[2] - bbox[0])
            original_height_emu = int(bbox_h_px * scale_y)
            bbox_w_emu = int(bbox_w_px * scale_x)
            inflation = (
                bbox_w_emu / column_width_emu
                if column_width_emu > 0 and bbox_w_emu > column_width_emu
                else 1.0
            )
            return int(original_height_emu * inflation * 1.3)
        return 914400  # 1 inch fallback

    # Text blocks
    if bbox:
        bbox_h_px = bbox[3] - bbox[1]
        bbox_w_px = max(1, bbox[2] - bbox[0])
        original_height_emu = int(bbox_h_px * scale_y)
        bbox_w_emu = int(bbox_w_px * scale_x)
        inflation = (
            bbox_w_emu / column_width_emu
            if column_width_emu > 0 and bbox_w_emu > column_width_emu
            else 1.0
        )
        return int(original_height_emu * inflation * LINE_HEIGHT_FACTOR)

    # No bbox — char-count based estimate
    font_size_emu = int(config.get("size", 12) * 12700)
    if column_width_emu > 0 and font_size_emu > 0:
        chars_per_line = max(1, column_width_emu / (font_size_emu * 0.52))
        num_lines = max(1, math.ceil(len(content) / chars_per_line))
    else:
        num_lines = max(1, len(content) // 80 + 1)
    return int(num_lines * font_size_emu * LINE_HEIGHT_FACTOR)


def _col_widths_emu_from_gaps(x_gaps, num_cols, usable_width_emu, page_width_px):
    """Compute per-column widths in EMU from _x_gaps gap data.

    Args:
        x_gaps: List of (gap_start_px, gap_end_px) tuples (column separator gaps).
        num_cols: Expected number of columns.
        usable_width_emu: Total usable page width in EMU.
        page_width_px: Original page width in pixels.

    Returns:
        List[int]: Column widths in EMU, length == num_cols.
        Falls back to equal-width split if gaps are absent or inconsistent.
    """
    if x_gaps and len(x_gaps) == num_cols - 1:
        col_edges = []
        prev_end = 0
        for gap_start, gap_end in x_gaps:
            col_edges.append((prev_end, gap_start))
            prev_end = gap_end + 1
        col_edges.append((prev_end, page_width_px))
        col_widths_px = [max(1, e - s) for s, e in col_edges]
        gap_widths_px = [g[1] - g[0] for g in x_gaps]
        total_px = sum(col_widths_px) + sum(gap_widths_px)
        if total_px > 0:
            px_to_emu = usable_width_emu / total_px
            return [int(w * px_to_emu) for w in col_widths_px]
    # Equal-width fallback
    col_w = usable_width_emu // max(num_cols, 1)
    return [col_w] * num_cols


def _minimize_section_break_para(para):
    """Minimize the height of a section-break paragraph.

    After doc.add_section(WD_SECTION.CONTINUOUS), the last paragraph becomes the
    section-break carrier. Set its font size and line spacing to 1pt so it contributes
    negligible vertical space to the page.

    Args:
        para: docx Paragraph object (typically doc.paragraphs[-1]).
    """
    from docx.shared import Emu as _Emu
    from docx.shared import Pt as _Pt

    para.paragraph_format.space_before = _Emu(0)
    para.paragraph_format.space_after = _Emu(0)
    if not para.runs:
        para.add_run()
    para.runs[0].font.size = _Pt(1)
    para.paragraph_format.line_spacing = _Pt(1)


def _estimate_page_content_height(
    segments, page_metrics, abs_image_paths, scale_y, original_image_width=0
):
    """Estimate total vertical content height for one page (EMU).

    Args:
        segments: List of segment dicts (type, blocks/columns).
        page_metrics: Dict from _build_page_metrics().
        abs_image_paths: Dict mapping image name to absolute path.
        scale_y: Pixels-to-EMU Y factor.
        original_image_width: Original page width in pixels; used as page_width_px
            for _col_widths_emu_from_gaps and passed to _estimate_block_height for
            accurate image aspect-ratio calculation.

    Returns:
        Tuple[int, Dict]: (estimated_height_emu, spacings_cache) where spacings_cache
            maps segment index to pre-computed spacing lists (single: list, multi:
            list-of-lists) so the write loop can reuse them without recomputing.
    """
    scale_x = page_metrics["scale_x"]
    usable_width_emu = page_metrics["usable_width_emu"]
    page_width_px = original_image_width if original_image_width > 0 else 1000
    total = 0
    spacings_cache: Dict[int, Any] = {}

    for seg_idx, segment in enumerate(segments):
        seg_type = segment["type"]

        if seg_type == "single":
            blocks = segment["blocks"]
            spacings = _compute_vertical_spacing(blocks, scale_y)
            spacings_cache[seg_idx] = spacings
            seg_height = 0
            for block, sp in zip(blocks, spacings):
                seg_height += _estimate_block_height(
                    block,
                    usable_width_emu,
                    abs_image_paths,
                    scale_x,
                    scale_y,
                    original_image_width,
                )
                if sp:
                    seg_height += sp
            total += seg_height
        else:
            # Multi-column: take the tallest column
            columns = segment["columns"]
            num_cols = len(columns)

            col_widths_emu = _col_widths_emu_from_gaps(
                segment.get("_x_gaps", []),
                num_cols,
                usable_width_emu,
                page_width_px,
            )

            col_spacings_list = []
            col_heights = []
            for col_idx, col_blocks in enumerate(columns):
                col_w = (
                    col_widths_emu[col_idx]
                    if col_idx < len(col_widths_emu)
                    else col_widths_emu[-1]
                )
                spacings = _compute_vertical_spacing(col_blocks, scale_y)
                col_spacings_list.append(spacings)
                ch = 0
                for block, sp in zip(col_blocks, spacings):
                    ch += _estimate_block_height(
                        block,
                        col_w,
                        abs_image_paths,
                        scale_x,
                        scale_y,
                        original_image_width,
                    )
                    if sp:
                        ch += sp
                col_heights.append(ch)
            spacings_cache[seg_idx] = col_spacings_list
            total += max(col_heights) if col_heights else 0

    # Section break overhead: each CONTINUOUS break ≈ 1 line (12pt ≈ 152400 EMU)
    section_break_count = max(0, len(segments) - 1)
    total += section_break_count * 152400

    return total, spacings_cache


def _compute_horizontal_indent(block, content_x1_px, page_width_px, scale_x):
    """Compute left_indent (EMU) for a single-column block.

    Only applies indent when the block's left edge is significantly offset
    from the content area's left edge (more than 3% of page width).
    Centered blocks (by config) are skipped.

    Args:
        block: Block dict with "bbox" and "config".
        content_x1_px: X coordinate of the content area left edge in pixels.
        page_width_px: Page width in pixels.
        scale_x: Pixels-to-EMU conversion factor for X axis.

    Returns:
        int or None: left_indent in EMU, or None for no indent.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    config = block.get("config") or {}
    if config.get("align") == WD_ALIGN_PARAGRAPH.CENTER:
        return None

    bbox = block.get("bbox")
    if not bbox:
        return None

    block_x1 = bbox[0]
    offset_px = block_x1 - content_x1_px
    threshold_px = page_width_px * 0.03
    if offset_px < threshold_px:
        return None

    indent_emu = int(offset_px * scale_x)
    indent_emu = min(indent_emu, 2743200)  # 3 inch cap
    return indent_emu if indent_emu > 0 else None


def _header_text_alignment(bbox, page_width_px):
    """Return tab prefix for header/footer companion text based on bbox position."""
    if not bbox or page_width_px <= 0:
        return "\t\t"  # fallback: right
    x_center = (bbox[0] + bbox[2]) / 2
    ratio = x_center / page_width_px
    if ratio < 0.35:
        return ""  # left — no tab
    elif ratio < 0.65:
        return "\t"  # center — one tab
    else:
        return "\t\t"  # right — two tabs


def _write_hf_image_block(
    section_hf,
    block,
    next_block,
    page_idx,
    abs_image_paths,
    original_image_width,
    usable_width_emu,
    companion_label,
):
    """Write header/footer image block with optional companion text merge.

    Args:
        section_hf: section.header or section.footer object.
        block: Current header_image/footer_image block dict.
        next_block: Next block in word_blocks (or None).
        page_idx: Current page index.
        abs_image_paths: Dict mapping image path to absolute path.
        original_image_width: Original page image width in pixels.
        usable_width_emu: Usable page width in EMU.
        companion_label: "header" or "footer" — label of the companion text block.

    Returns:
        True if next_block was consumed (caller should skip it), False otherwise.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    content = block.get("content", "").strip()
    abs_path = abs_image_paths.get(content)
    if not abs_path:
        return False

    next_same_page = (
        next_block is not None
        and next_block.get("type") == companion_label
        and next_block.get("page_index", 0) == page_idx
    )
    para = section_hf.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    img_w = _header_image_width(
        abs_path,
        block.get("bbox"),
        original_image_width,
        usable_width_emu,
    )
    para.add_run().add_picture(abs_path, width=img_w)
    consumed_next = False
    if next_same_page:
        next_content = next_block.get("content", "").strip()
        if next_content:
            tab_prefix = _header_text_alignment(
                next_block.get("bbox"), original_image_width
            )
            run = para.add_run(tab_prefix + next_content)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(9)
        consumed_next = True
    return consumed_next


class WordConverter:
    """Convert structured word_blocks to a :class:`docx.Document`."""

    @staticmethod
    def convert(
        word_blocks: List[Dict],
        *,
        abs_image_paths: Dict[str, str],
        original_image_width: int = 500,
        original_image_height: int = 0,
    ):
        """Convert word_blocks to a docx.Document object.

        Args:
            word_blocks: List[Dict] — each dict has keys "type", "content", "config",
                optional "page_index".
            abs_image_paths: Dict[str, str] — {original_path: abs_path} from save_images().
            original_image_width: int — reserved for future scaling (currently unused).
            original_image_height: int — used to detect landscape orientation.

        Returns:
            docx.Document
        """
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Emu

        # Detect landscape: width must exceed height by at least 20%
        is_landscape = (
            original_image_height > 0
            and original_image_width > original_image_height * 1.2
        )

        # A4 portrait:  210mm × 297mm  →  7560820 × 10693400 EMU
        # A4 landscape: 297mm × 210mm  → 10693400 × 7560820 EMU
        _PAGE_W = 10693400 if is_landscape else 7560820
        _PAGE_H = 7560820 if is_landscape else 10693400
        # Usable width assuming default 1-inch margins (914400 EMU each side)
        _USABLE_W = _PAGE_W - 2 * 914400

        def _apply_page_size(section):
            if is_landscape:
                section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Emu(_PAGE_W)
            section.page_height = Emu(_PAGE_H)

        doc = Document()
        _apply_page_size(doc.sections[0])
        current_page = None
        consumed: set = set()

        for i, block in enumerate(word_blocks):
            if i in consumed:
                continue

            page_idx = block.get("page_index", 0)
            if current_page is None:
                current_page = page_idx
            elif page_idx != current_page:
                _apply_page_size(doc.add_section())
                current_page = page_idx

            label = block.get("type")
            content = block.get("content", "").strip()
            config = block.get("config") or {}

            # --- header/footer ---
            if label == "header" and content:
                section = doc.sections[-1]
                section.header.is_linked_to_previous = False
                para = section.header.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(content)
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            elif label == "footer" and content:
                section = doc.sections[-1]
                section.footer.is_linked_to_previous = False
                para = section.footer.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(content)
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            elif label == "header_image" and content:
                if abs_image_paths.get(content):
                    section = doc.sections[-1]
                    section.header.is_linked_to_previous = False
                    next_block = (
                        word_blocks[i + 1] if i + 1 < len(word_blocks) else None
                    )
                    if _write_hf_image_block(
                        section.header,
                        block,
                        next_block,
                        page_idx,
                        abs_image_paths,
                        original_image_width,
                        _USABLE_W,
                        "header",
                    ):
                        consumed.add(i + 1)
            elif label == "footer_image" and content:
                if abs_image_paths.get(content):
                    section = doc.sections[-1]
                    section.footer.is_linked_to_previous = False
                    next_block = (
                        word_blocks[i + 1] if i + 1 < len(word_blocks) else None
                    )
                    if _write_hf_image_block(
                        section.footer,
                        block,
                        next_block,
                        page_idx,
                        abs_image_paths,
                        original_image_width,
                        _USABLE_W,
                        "footer",
                    ):
                        consumed.add(i + 1)
            else:
                _write_block(
                    doc,
                    block,
                    abs_image_paths,
                    original_image_width,
                    usable_width_emu=_USABLE_W,
                )

        return doc
